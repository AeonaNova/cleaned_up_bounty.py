import tkinter as tk
from tkinter import IntVar, Checkbutton, Button
from typing import List, Any
from docx import Document
from datetime import datetime
n: datetime = datetime.now()
d: str = n.strftime("%d")
m: str = n.strftime("%m")
y: str = n.strftime("%y")

with open("addresses.txt", 'r', encoding="utf-8") as file:
    addresses: list[str] = file.read().split('\n')

# Creating document .docx
doc: Any = Document()

def save_selections(counter:int = 0) -> None:
    """
    if address pointed in checkbox tkinter gui, then that address placed in list and added in docx
    """
    selected_addresses: list[str] = []
    s_var: IntVar
    j: int
    for j, s_var in enumerate(sel_vars):
        if s_var.get():
            selected_addresses.append(addresses[j])
            counter += 1

    # Add selected addresses in document
    first = f"На подведомственной территории района Теплый Стан по состоянию на {d}.{m}.{y} г. зачищено {counter} дворовых территорий:"
    doc.add_paragraph(first)
    sel_address: str
    for sel_address in selected_addresses:
        doc.add_paragraph(sel_address)

    doc.save("убранные_адреса.docx")


# Creating gui with checkboxes
root:tk = tk.Tk()
root.title("Выбор адресов")

sel_vars: list[IntVar] = []
i: int
address: str
for i, address in enumerate(addresses):
    var: IntVar = tk.IntVar()
    checkbox: Checkbutton = tk.Checkbutton(root, text=address, variable=var)
    checkbox.grid(row=i%40, column=i//40, sticky='w')
    sel_vars.append(var)

save_button: Button = tk.Button(root, text="Сохранить", command=save_selections)
save_button.grid(row=(len(addresses)+1)//3, columnspan=3)

root.mainloop()
